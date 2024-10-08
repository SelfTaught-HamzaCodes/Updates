# Model for application, Data Storage and Communications with API.
import textract
import re
import os
import sys
import pandas as pd

# Imports to Generate Label Document:
from docx import Document
from docxtpl import DocxTemplate
from lxml import etree


class Model:
    def __init__(self):

        # Initialise the paths for Excel (Configurations) and Word Files:
        self.excel_path = None
        self.excel_sheet = None
        self.excel_row = None

        self.word_path = None
        self.placeholders = []


    # Get abs_path to source:
    @staticmethod
    def resource_path(relative_path):
        """ Get absolute path to resource, works for dev and for PyInstaller """
        try:
            # PyInstaller creates a temp folder and stores path in _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)

    # Set values for Excel file and Configurations:
    def set_excel_values(self, excel, sheet, row):

        self.excel_path = excel
        self.excel_sheet = sheet
        self.excel_row = row

    # Get values for Excel file and Configurations:
    def get_excel_values(self):

        return self.excel_path, self.excel_sheet, self.excel_row

    # Get Column Names:
    def get_column_names(self):

        # Generate Dataframe:
        excel_dataframe = pd.read_excel(self.excel_path, sheet_name=self.excel_sheet, header=int(self.excel_row) - 1)

        return list(excel_dataframe.columns)

    # Set path for Word file:
    def set_word_file(self, word_file):

        # Set File:
        self.word_path = word_file
    
        # Generate PlaceHolders:
        self.set_placeholders()

    # Get path for Word file:
    def get_word_file(self):

        return self.word_path


    # Set Placeholders:
    def set_placeholders(self):

        # Process Word File:
        processed = textract.process(self.word_path)
        processed = processed.decode("utf-8")

        # Extract Placeholders:
        expression = re.compile(r"{{(.*)}}")

        # Store Placeholders:
        self.placeholders = expression.findall(processed)

        # Sort Placeholders:
        self.placeholders.sort()

    # Extract Placeholders:
    def get_placeholders(self):

        return self.placeholders


    # Generate Label Document:
    def generate_label_doc(self, elements):

        # Open Excel File:
        packing_list = pd.read_excel(self.excel_path, sheet_name=self.excel_sheet, header=int(self.excel_row) - 1)

        # Open Label File:
        doc = DocxTemplate(self.word_path)

        # Create a combined document to store all labels
        final_doc = Document()

        # Generate Labels:
        for i in range(len(packing_list)):

            context = {}

            for placeholder, element in elements.items():

                # If Dropdown had a value:
                if element["dropdown_element"].value:
                    value = packing_list[element["dropdown_element"].value].iloc[i]
                    if isinstance(value, (int, float)):
                        # Convert the value to a float, round to 3 decimal places, and format to retain trailing zeros
                        context[placeholder] = f"{float(value):.3f}"
                    else:
                        context[placeholder] = str(value)
                    continue

                elif element["dropdown_element"].value:
                    context[placeholder] = element["dropdown_element"].value
                    continue

                # Add Placeholder as the value, if not mentioned:
                context[placeholder] = placeholder

            # Render the document
            doc.render(context)

            # Save the rendered document to a temporary file
            temp_file = f'temp_rendered_doc_{i}.docx'
            doc.save(temp_file)

            # Open the temporary rendered document
            rendered_doc = Document(temp_file)

            # Extract the table from the rendered document
            rendered_table = rendered_doc.tables[0]  # Adjust the table index as needed

            # Create a new table in the final document
            final_table = final_doc.add_table(rows=len(rendered_table.rows), cols=len(rendered_table.columns))

            # Get the XML element of the rendered table
            table_xml = rendered_table._tbl.xml

            # Create a new table element in the final document
            new_table = etree.fromstring(table_xml)

            # Append the new table element to the final table
            final_table._tbl.append(new_table)

            # Add a page break after each rendered document
            final_doc.add_page_break()

            # Remove the temporary rendered document file
            # Uncomment the line below if you want to remove the temporary files
            os.remove(temp_file)

        # Return this and save it at designated spot:
        return final_doc
